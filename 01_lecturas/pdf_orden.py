"""Genera orden_verificacion_YYYY-MM.pdf para el operario.

Solo incluye las anomalías que requieren ir al predio:
    SIN_LECTURA · MEDIDOR_INVERTIDO · POSIBLE_CAMBIO_MEDIDOR · EXCESIVO

Genera primero un .docx con python-docx y luego lo convierte a PDF con docx2pdf.
Si Word no está disponible, deja el .docx y avisa.
"""
from __future__ import annotations

import logging
import sys
from collections import Counter
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, RGBColor, Cm, Inches

sys.path.insert(0, str(Path(__file__).parent))
import config

log = logging.getLogger(__name__)

ANOMALIAS_DE_CAMPO = {"SIN_LECTURA", "MEDIDOR_INVERTIDO", "POSIBLE_CAMBIO_MEDIDOR", "EXCESIVO"}

# Color por tipo (RGB hex sin #)
COLOR_HEAD = {
    "SIN_LECTURA":            "3730A3",
    "MEDIDOR_INVERTIDO":      "7F1D1D",
    "POSIBLE_CAMBIO_MEDIDOR": "5B21B6",
    "EXCESIVO":               "9A3412",
}

# Mensajes legibles para el operario (sin jerga técnica)
def _mensaje_para_operario(tipo: str, ant=None, act=None, m3=None) -> str:
    if tipo == "SIN_LECTURA":
        return ("No se registró la lectura este mes. El operario no anotó el valor del medidor, "
                "o el predio estaba cerrado al momento de la visita.")
    if tipo == "MEDIDOR_INVERTIDO":
        return (f"El medidor marca un poco menos que el mes pasado. Lectura actual {act:.0f} vs "
                f"{ant:.0f} del mes pasado · diferencia de {act - ant:.0f} m³. "
                f"Es posible que el medidor esté instalado al revés.")
    if tipo == "POSIBLE_CAMBIO_MEDIDOR":
        return (f"El medidor marca mucho menos que el mes pasado. Lectura actual {act:.0f} vs "
                f"{ant:.0f} del mes pasado · diferencia de {act - ant:.0f} m³. "
                f"Probablemente se cambió el medidor sin reportar.")
    if tipo == "EXCESIVO":
        consumo = act - ant
        return (f"Consumo muy alto. Este predio registró {consumo:.0f} m³ este mes · "
                f"el umbral normal es {int(config.M3_EXCESIVO)} m³. Puede haber fuga.")
    return ""


CAUSAS = {
    "SIN_LECTURA":            ["Predio cerrado", "Medidor sin display visible", "Omisión del operario"],
    "MEDIDOR_INVERTIDO":      ["Medidor instalado al revés", "Lectura mal anotada", "Medidor en mal estado"],
    "POSIBLE_CAMBIO_MEDIDOR": ["Medidor cambiado y no reportado", "Lectura tomada en medidor equivocado",
                              "Error de transcripción grave"],
    "EXCESIVO":               ["Fuga visible o en cisterna", "Medidor en mal estado",
                              "Uso intensivo legítimo", "Lectura mal anotada"],
}

ACCIONES = {
    "SIN_LECTURA":            ("Ir al predio, anotar el número del medidor abajo. "
                              "Si el predio sigue cerrado, escribir P en obs."),
    "MEDIDOR_INVERTIDO":      ("Ir al predio y verificar si el medidor está instalado al revés. "
                              "Anotar la lectura correcta. Si se confirma medidor cambiado, escribir M en obs."),
    "POSIBLE_CAMBIO_MEDIDOR": ("Ir al predio, confirmar con el usuario si se cambió el medidor. "
                              "Anotar la lectura del nuevo medidor. Si se cambió, escribir M en obs."),
    "EXCESIVO":               ("Ir al predio, revisar si hay fuga visible (caños, tanque, cisterna). "
                              "Confirmar el número del medidor. Si hay fuga, escribir F en obs."),
}


def _set_cell_bg(cell, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _add_cover(doc: Document, mes_ano: str, conteo: Counter) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("JASS — Orden de Verificación")
    run.bold = True
    run.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Ciclo {mes_ano}")
    run.font.size = Pt(14)
    doc.add_paragraph()

    # Datos
    metadata = [
        ("Emitido:", date.today().strftime("%d/%m/%Y")),
        ("Operario:", ""),
        ("Casos a verificar:", str(sum(conteo.values()))),
        ("Fecha de devolución:", ""),
    ]
    tabla = doc.add_table(rows=len(metadata), cols=2)
    tabla.columns[0].width = Cm(5)
    tabla.columns[1].width = Cm(10)
    for i, (k, v) in enumerate(metadata):
        c1 = tabla.cell(i, 0)
        c1.text = k
        c1.paragraphs[0].runs[0].bold = True
        tabla.cell(i, 1).text = v
    doc.add_paragraph()

    # Resumen
    p = doc.add_paragraph()
    run = p.add_run("Resumen de casos")
    run.bold = True
    run.font.size = Pt(13)

    res = doc.add_table(rows=len(conteo) + 1, cols=3)
    for i, (tipo, n) in enumerate(conteo.most_common()):
        c0 = res.cell(i, 0)
        c0.text = tipo
        c0.paragraphs[0].runs[0].bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(11)
        res.cell(i, 1).text = str(n)
        accion_corta = {
            "SIN_LECTURA":            "→ ir al predio y leer",
            "MEDIDOR_INVERTIDO":      "→ verificar si el medidor está al revés",
            "POSIBLE_CAMBIO_MEDIDOR": "→ confirmar si cambiaron el medidor",
            "EXCESIVO":               "→ revisar si hay fuga",
        }.get(tipo, "")
        res.cell(i, 2).text = accion_corta
    res.cell(len(conteo), 0).text = "TOTAL"
    res.cell(len(conteo), 0).paragraphs[0].runs[0].bold = True
    res.cell(len(conteo), 1).text = str(sum(conteo.values()))
    res.cell(len(conteo), 1).paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Instrucciones para el operario:")
    run.bold = True
    doc.add_paragraph("1. Cada caso ocupa una página · llevar el documento al predio.")
    doc.add_paragraph("2. Confirmar el número del medidor y leer la marcación actual.")
    doc.add_paragraph("3. Anotar el valor en el recuadro de cada caso (escribir con lapicero).")
    doc.add_paragraph("4. Códigos obs: M = medidor cambiado · F = fuga visible · P = predio cerrado.")
    doc.add_paragraph("5. Devolver el documento firmado al supervisor.")


def _add_caso(doc: Document, b: dict) -> None:
    doc.add_page_break()

    tipo = b["tipo"]
    color = COLOR_HEAD.get(tipo, "1E3A5F")

    # Header del caso: una tabla de 1 fila con color de fondo
    head = doc.add_table(rows=1, cols=2)
    cell_l = head.cell(0, 0)
    cell_r = head.cell(0, 1)
    _set_cell_bg(cell_l, color)
    _set_cell_bg(cell_r, color)

    pl = cell_l.paragraphs[0]
    run = pl.add_run(f"{b['mz']}-{b['lt']}")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("FFFFFF")
    run.font.size = Pt(18)

    pl2 = cell_l.add_paragraph()
    run = pl2.add_run(b.get("nombre", ""))
    run.font.color.rgb = RGBColor.from_string("FFFFFF")
    run.font.size = Pt(12)

    pr = cell_r.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = pr.add_run(tipo.replace("_", " "))
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("FFFFFF")
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Datos del sistema
    p = doc.add_paragraph()
    p.add_run("¿Qué dice el sistema?").bold = True

    marc_ant = b.get("marc_ant_hist") if b.get("marc_ant_hist") is not None else b.get("marc_ant", "")
    marc_act = b.get("marc_act", "") or "— sin dato —"
    m3       = b.get("m3", "") or "— sin dato —"
    fecha    = date.today().strftime("%d/%m/%Y")

    datos = doc.add_table(rows=2, cols=4)
    for i, label in enumerate(["Lectura mes pasado", "Lectura actual registrada", "M3 registrado", "Fecha"]):
        c = datos.cell(0, i)
        c.text = label
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0].bold = True
    for i, val in enumerate([str(marc_ant), str(marc_act), str(m3), fecha]):
        c = datos.cell(1, i)
        c.text = val
        c.paragraphs[0].runs[0].font.size = Pt(12)

    doc.add_paragraph()

    # Mensaje
    p = doc.add_paragraph()
    p.add_run("¿Qué pasa con este predio?").bold = True

    msg_kw = {}
    if b.get("marc_ant_hist") is not None and b.get("marc_act"):
        try:
            msg_kw["ant"] = float(b["marc_ant_hist"])
            msg_kw["act"] = float(b["marc_act"])
        except (ValueError, TypeError):
            pass
    mensaje = _mensaje_para_operario(tipo, **msg_kw)
    p = doc.add_paragraph(mensaje)
    p.paragraph_format.left_indent = Cm(0.5)

    # Causas
    p = doc.add_paragraph()
    p.add_run("Causas posibles: ").bold = True
    for causa in CAUSAS.get(tipo, []):
        p.add_run(f"• {causa}   ")

    # Acción
    p = doc.add_paragraph()
    p.add_run("Acción: ").bold = True
    p.add_run(ACCIONES.get(tipo, ""))

    doc.add_paragraph()

    # Espacios para escribir
    p = doc.add_paragraph()
    run = p.add_run("✏️ Qué encontraste en campo")
    run.bold = True
    run.font.size = Pt(11)

    campos = [
        ("Lectura verificada:", ""),
        ("M3 calculado:", "(lectura verificada − " + str(marc_ant) + ")"),
        ("obs (M / F / P / —):", ""),
        ("Notas:", ""),
    ]
    for label, hint in campos:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run("  " + "_" * 40)
        if hint:
            p2 = doc.add_paragraph()
            run = p2.add_run(hint)
            run.font.size = Pt(9)
            run.italic = True
            p2.paragraph_format.left_indent = Cm(3)


def _add_firmas(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph("_" * 30 + "        " + "_" * 30)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Firma operario             Firma supervisor (recibió)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Fecha: __ / __ / ____          Fecha: __ / __ / ____")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def generar_pdf(bloqueantes: list[dict], mes_ano: str) -> Path | None:
    """Genera orden_verificacion_YYYY-MM.pdf con los bloqueantes de campo.

    Devuelve la ruta del PDF generado, o None si no hay casos de campo.
    """
    casos = [b for b in bloqueantes if b["tipo"] in ANOMALIAS_DE_CAMPO]
    if not casos:
        log.info("No hay anomalías de campo · no se genera PDF")
        return None

    # Ordenar por tipo y luego MZ-LT (SIN_LECTURA primero, EXCESIVO último)
    orden = {"SIN_LECTURA": 0, "MEDIDOR_INVERTIDO": 1, "POSIBLE_CAMBIO_MEDIDOR": 2, "EXCESIVO": 3}
    casos.sort(key=lambda b: (orden.get(b["tipo"], 9), b["mz"], b["lt"]))

    conteo = Counter(b["tipo"] for b in casos)

    config.OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = config.OUTPUTS_DIR / f"orden_verificacion_{mes_ano}.docx"
    pdf_path = config.orden_verificacion_path(mes_ano)

    doc = Document()
    # Márgenes A4
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    _add_cover(doc, mes_ano, conteo)
    for b in casos:
        _add_caso(doc, b)
    _add_firmas(doc)

    doc.save(docx_path)
    log.info(f"orden_verificacion_{mes_ano}.docx generado · {len(casos)} casos de campo")

    # Convertir a PDF
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        # Quitar el .docx intermedio (opcional, lo dejo por ahora)
        log.info(f"orden_verificacion_{mes_ano}.pdf generado")
        return pdf_path
    except ImportError:
        log.warning("docx2pdf no está instalado · se generó solo el .docx")
        return docx_path
    except Exception as e:
        log.warning(f"No se pudo convertir a PDF ({e}) · se generó solo el .docx")
        return docx_path
